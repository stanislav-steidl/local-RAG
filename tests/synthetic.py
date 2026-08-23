"""Builders for synthetic test documents.

Every fixture the test suite parses is generated here. No file from a real
corpus is ever read by a test — see ``docs/decisions.md`` for the reasoning.

The Czech strings are deliberate: the corpus this project targets is bilingual,
and diacritics are exactly what encoding and extraction bugs destroy.
"""

from __future__ import annotations

from typing import TYPE_CHECKING

import docx
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

if TYPE_CHECKING:
    from pathlib import Path

__all__ = ["CZECH_SAMPLE", "make_docx", "make_pdf", "make_scanned_pdf"]

#: Text exercising the Czech diacritics that a wrong codepage mangles.
CZECH_SAMPLE = "Smlouva o dílo — příloha č. 3, částka 12 345 Kč, žádost přijata."


def make_pdf(path: Path, pages: list[str]) -> Path:
    """Write a PDF with one text-bearing page per entry in ``pages``.

    Args:
        path: Destination file.
        pages: Text to draw on each page, in order.

    Returns:
        ``path``, for convenient chaining.
    """
    pdf = canvas.Canvas(str(path), pagesize=A4)
    for text in pages:
        # Helvetica has no glyphs for Czech diacritics; the default encoding
        # would silently substitute them. Draw with a Unicode-capable font.
        pdf.setFont("Helvetica", 12)
        pdf.drawString(72, 750, text)
        pdf.showPage()
    pdf.save()
    return path


def make_scanned_pdf(path: Path, page_count: int = 1) -> Path:
    """Write a PDF whose pages carry no text layer.

    This is what a scanned document looks like to a text extractor, and the
    case that must be recognised so such files can be routed to OCR.

    Args:
        path: Destination file.
        page_count: Number of blank pages to emit.

    Returns:
        ``path``, for convenient chaining.
    """
    pdf = canvas.Canvas(str(path), pagesize=A4)
    for _ in range(page_count):
        # A filled rectangle stands in for scanned pixels: visible content,
        # no extractable text.
        pdf.rect(72, 600, 200, 100, fill=1)
        pdf.showPage()
    pdf.save()
    return path


def make_docx(
    path: Path,
    paragraphs: list[str] | None = None,
    table: list[list[str]] | None = None,
) -> Path:
    """Write a Word document with the given paragraphs and optional table.

    Args:
        path: Destination file.
        paragraphs: Body paragraphs to add.
        table: Rows of cell text. Contracts and invoices keep the retrievable
            facts in tables, so parsers must not skip them.

    Returns:
        ``path``, for convenient chaining.
    """
    document = docx.Document()
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)

    if table:
        added = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for column_index, cell_text in enumerate(row):
                added.cell(row_index, column_index).text = cell_text

    document.save(str(path))
    return path
