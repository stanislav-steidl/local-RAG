"""Concrete parsers for the formats the corpus actually contains.

Each parser imports its third-party dependency lazily. Those libraries live
behind the ``parsing`` extra, so :mod:`local_rag.ingest` must stay importable
without them — a missing dependency should surface as an actionable error when
a PDF is parsed, not as an ImportError when the package is first imported.
"""

from __future__ import annotations

import logging
from typing import TYPE_CHECKING, ClassVar

from local_rag.ingest.base import DocumentParseError, DocumentParser, MissingDependencyError
from local_rag.models import PageText

if TYPE_CHECKING:
    from pathlib import Path

__all__ = ["DocxParser", "PdfParser", "TextParser"]

logger = logging.getLogger(__name__)

#: Encodings tried in order when reading plain text. UTF-8 covers modern files;
#: cp1250 is the Windows Central European codepage that older Czech documents
#: were written in. latin-1 is last because it decodes any byte sequence
#: without raising, so it must never pre-empt a more accurate candidate.
_TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1250", "latin-1")


class TextParser(DocumentParser):
    """Reads plain-text formats, tolerating legacy encodings."""

    extensions: ClassVar[frozenset[str]] = frozenset({".txt", ".md", ".csv"})
    is_paginated: ClassVar[bool] = False

    def parse(self, path: Path) -> tuple[PageText, ...]:
        """Read ``path`` as text, trying each candidate encoding in turn.

        Args:
            path: File to read.

        Returns:
            A single page holding the file's text, or no pages if the file is
            empty.

        Raises:
            DocumentParseError: If the file cannot be read from disk.
        """
        try:
            raw = path.read_bytes()
        except OSError as error:
            raise DocumentParseError(f"cannot read {path}: {error}") from error

        if not raw:
            return ()

        for encoding in _TEXT_ENCODINGS:
            try:
                text = raw.decode(encoding)
            except UnicodeDecodeError:
                continue
            if encoding != _TEXT_ENCODINGS[0]:
                logger.debug("Decoded %s as %s", path, encoding)
            return (PageText(page_number=1, text=text),)

        # Unreachable while latin-1 remains last: it decodes any byte sequence.
        # Kept so that removing it from _TEXT_ENCODINGS fails loudly instead of
        # silently returning no pages.
        raise DocumentParseError(f"no candidate encoding could decode {path}")  # pragma: no cover


class PdfParser(DocumentParser):
    """Extracts the text layer of a PDF, one page at a time."""

    extensions: ClassVar[frozenset[str]] = frozenset({".pdf"})
    is_paginated: ClassVar[bool] = True

    def parse(self, path: Path) -> tuple[PageText, ...]:
        """Extract each page's text layer.

        Pages whose text layer is empty — the signature of a scanned document —
        are returned as empty pages rather than dropped, so that page numbering
        continues to match the physical document and the caller can see that
        extraction found nothing.

        Args:
            path: PDF to read.

        Returns:
            One page per page of the PDF.

        Raises:
            DocumentParseError: If the file is not a readable PDF.
            MissingDependencyError: If pdfplumber is not installed.
        """
        try:
            import pdfplumber  # noqa: PLC0415  # optional dependency, imported on use
        except ImportError as error:
            raise MissingDependencyError(
                "reading PDFs requires pdfplumber: pip install 'local-rag[parsing]'"
            ) from error

        try:
            with pdfplumber.open(path) as pdf:
                return tuple(
                    PageText(page_number=number, text=page.extract_text() or "")
                    for number, page in enumerate(pdf.pages, start=1)
                )
        except Exception as error:  # pdfplumber raises a variety of low-level errors
            raise DocumentParseError(f"cannot parse PDF {path}: {error}") from error


class DocxParser(DocumentParser):
    """Extracts paragraph and table text from a Word document.

    Table cells are included because contracts and invoices routinely put the
    facts worth retrieving — dates, amounts, parties — inside tables.
    """

    extensions: ClassVar[frozenset[str]] = frozenset({".docx"})
    is_paginated: ClassVar[bool] = False

    def parse(self, path: Path) -> tuple[PageText, ...]:
        """Extract body text and table contents.

        DOCX page breaks are decided at render time, so the whole document is
        returned as a single page.

        Args:
            path: Word document to read.

        Returns:
            A single page holding the document's text, or no pages if it holds
            no text at all.

        Raises:
            DocumentParseError: If the file is not a readable DOCX.
            MissingDependencyError: If python-docx is not installed.
        """
        try:
            import docx  # noqa: PLC0415  # optional dependency, imported on use
        except ImportError as error:
            raise MissingDependencyError(
                "reading DOCX files requires python-docx: pip install 'local-rag[parsing]'"
            ) from error

        try:
            document = docx.Document(str(path))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            blocks.extend(
                cell.text for table in document.tables for row in table.rows for cell in row.cells
            )
        except Exception as error:  # python-docx raises package-specific errors
            raise DocumentParseError(f"cannot parse DOCX {path}: {error}") from error

        text = "\n".join(block for block in blocks if block.strip())
        if not text:
            return ()
        return (PageText(page_number=1, text=text),)
